"""
docx_builder.py — Builds a DOCX resume from the user's template + optimized content.
No LLM involvement. Pure local DOCX assembly.

Strategy: Copy the template, clear content paragraphs/tables, then re-populate
with optimized content while preserving all style formatting (fonts, colors,
borders, spacing, alignment).
"""

import io
import copy
import json
import re
from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────
# Style Snapshot — captures formatting from a run
# ─────────────────────────────────────────────
def _snapshot_run_format(run):
    """Take a snapshot of a run's formatting for later reuse."""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        return copy.deepcopy(rPr)
    return None


def _apply_run_format(run, rPr_snapshot):
    """Apply a previously-snapshotted format to a run."""
    if rPr_snapshot is not None:
        existing = run._element.find(qn('w:rPr'))
        if existing is not None:
            run._element.remove(existing)
        run._element.insert(0, copy.deepcopy(rPr_snapshot))


def _snapshot_para_format(para):
    """Take a snapshot of a paragraph's pPr (alignment, borders, spacing, etc.)."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is not None:
        return copy.deepcopy(pPr)
    return None


def _apply_para_format(para, pPr_snapshot):
    """Apply a previously-snapshotted paragraph format."""
    if pPr_snapshot is not None:
        existing = para._element.find(qn('w:pPr'))
        if existing is not None:
            para._element.remove(existing)
        para._element.insert(0, copy.deepcopy(pPr_snapshot))


# ─────────────────────────────────────────────
# Template Analyzer — learns the template's styles
# ─────────────────────────────────────────────
class TemplateStyles:
    """Extracts and stores style snapshots from the reference template."""

    def __init__(self, template_path: str):
        self.doc = Document(template_path)
        self.template_path = template_path
        self._analyze()

    def _analyze(self):
        """Walk the template and capture formatting for each element type."""
        paras = self.doc.paragraphs

        # --- Header styles (paragraphs 0, 1, 2) ---
        self.name_para_fmt = _snapshot_para_format(paras[0])
        self.name_run_fmt = _snapshot_run_format(paras[0].runs[0]) if paras[0].runs else None

        self.subtitle_para_fmt = _snapshot_para_format(paras[1])
        self.subtitle_run_fmt = _snapshot_run_format(paras[1].runs[0]) if paras[1].runs else None

        self.contact_para_fmt = _snapshot_para_format(paras[2])
        # Contact has two runs: normal text + links (different color)
        self.contact_run_fmt = _snapshot_run_format(paras[2].runs[0]) if paras[2].runs else None
        self.contact_link_run_fmt = _snapshot_run_format(paras[2].runs[1]) if len(paras[2].runs) > 1 else self.contact_run_fmt

        # --- Section heading style (e.g. "PROFESSIONAL SUMMARY") ---
        # para 4 is the first section heading
        self.section_heading_para_fmt = _snapshot_para_format(paras[4])
        self.section_heading_run_fmt = _snapshot_run_format(paras[4].runs[0]) if paras[4].runs else None

        # --- Body text style (e.g. summary paragraph) ---
        self.body_para_fmt = _snapshot_para_format(paras[5])
        self.body_run_fmt = _snapshot_run_format(paras[5].runs[0]) if paras[5].runs else None

        # --- Job title line: "Data & GenAI Engineer   Jan 2024 – Present" ---
        # para 8: bold title run + gray dates run
        self.job_title_para_fmt = _snapshot_para_format(paras[8])
        self.job_title_run_fmt = _snapshot_run_format(paras[8].runs[0]) if paras[8].runs else None
        self.job_dates_run_fmt = _snapshot_run_format(paras[8].runs[1]) if len(paras[8].runs) > 1 else None

        # --- Company line: "Company Name  |  Location" ---
        # para 9: bold company + gray location
        self.company_para_fmt = _snapshot_para_format(paras[9])
        self.company_run_fmt = _snapshot_run_format(paras[9].runs[0]) if paras[9].runs else None
        self.company_loc_run_fmt = _snapshot_run_format(paras[9].runs[1]) if len(paras[9].runs) > 1 else None

        # --- Bullet point style (List Paragraph) ---
        # paras 10-14 are bullet points
        self.bullet_para_fmt = _snapshot_para_format(paras[10])
        self.bullet_run_fmt = _snapshot_run_format(paras[10].runs[0]) if paras[10].runs else None
        # Also capture the style name for numId/list reference
        self.bullet_style = paras[10].style

        # --- Project title line: "Project Name  | Tech Stack" ---
        # para 16: bold project name + gray tech stack
        self.project_title_para_fmt = _snapshot_para_format(paras[16])
        self.project_title_run_fmt = _snapshot_run_format(paras[16].runs[0]) if paras[16].runs else None
        self.project_tech_run_fmt = _snapshot_run_format(paras[16].runs[1]) if len(paras[16].runs) > 1 else None

        # --- Education line ---
        # para 26: bold degree + gray dates
        self.edu_para_fmt = _snapshot_para_format(paras[26])
        self.edu_run_fmt = _snapshot_run_format(paras[26].runs[0]) if paras[26].runs else None
        self.edu_dates_run_fmt = _snapshot_run_format(paras[26].runs[1]) if len(paras[26].runs) > 1 else None

        # --- Coursework line (italic) ---
        self.coursework_para_fmt = _snapshot_para_format(paras[27])
        self.coursework_run_fmt = _snapshot_run_format(paras[27].runs[0]) if paras[27].runs else None

        # --- Skills table styles ---
        table = self.doc.tables[0]
        # Category cell (left column) - first cell para
        cat_para = table.rows[0].cells[0].paragraphs[0]
        self.skill_cat_para_fmt = _snapshot_para_format(cat_para)
        self.skill_cat_run_fmt = _snapshot_run_format(cat_para.runs[0]) if cat_para.runs else None

        # Skills cell (right column) - first cell para
        skill_para = table.rows[0].cells[1].paragraphs[0]
        self.skill_val_para_fmt = _snapshot_para_format(skill_para)
        self.skill_val_run_fmt = _snapshot_run_format(skill_para.runs[0]) if skill_para.runs else None

        # Capture the table XML for structure (borders, column widths)
        self.table_xml = copy.deepcopy(self.doc.tables[0]._element)

        # --- Spacer paragraph (empty para between header and content) ---
        self.spacer_para_fmt = _snapshot_para_format(paras[3])


# ─────────────────────────────────────────────
# DOCX Builder — assembles the final resume
# ─────────────────────────────────────────────
def build_resume(template_path: str, optimized_data: dict) -> io.BytesIO:
    """Build a DOCX resume using the template's styling and optimized content.
    
    Args:
        template_path: Path to the reference DOCX template.
        optimized_data: Dict with keys:
            - name, subtitle, contact_line
            - summary
            - skills: list of {"category": str, "items": str}
            - experience: list of {"title": str, "dates": str, "company": str, "location": str, "bullets": list[str]}
            - projects: list of {"name": str, "tech": str, "bullets": list[str]}
            - education: list of {"degree": str, "institution": str, "dates": str, "coursework": str (optional)}
            - certifications: list[str]
            - additional: list[str]
    
    Returns:
        BytesIO containing the DOCX file.
    """
    styles = TemplateStyles(template_path)
    
    # Start from the template to preserve document-level settings
    # (page size, margins, default styles, numbering definitions)
    doc = Document(template_path)
    
    # Clear ALL existing content from the body
    body = doc.element.body
    # Remove all children except sectPr (page layout)
    children_to_remove = []
    for child in body:
        if child.tag != qn('w:sectPr'):
            children_to_remove.append(child)
    for child in children_to_remove:
        body.remove(child)

    # ── Helper functions ──
    def add_para(text, para_fmt, run_fmt, style=None):
        """Add a paragraph with specific formatting."""
        p = doc.add_paragraph()
        if style:
            p.style = style
        _apply_para_format(p, para_fmt)
        run = p.add_run(text)
        _apply_run_format(run, run_fmt)
        return p

    def add_multi_run_para(runs_data, para_fmt):
        """Add a paragraph with multiple differently-formatted runs.
        runs_data: list of (text, run_fmt) tuples.
        """
        p = doc.add_paragraph()
        _apply_para_format(p, para_fmt)
        for text, rfmt in runs_data:
            run = p.add_run(text)
            _apply_run_format(run, rfmt)
        return p

    def add_section_heading(text):
        """Add a section heading with border styling."""
        return add_para(text, styles.section_heading_para_fmt, styles.section_heading_run_fmt)

    def add_bullet(text):
        """Add a bullet point in List Paragraph style."""
        p = doc.add_paragraph()
        p.style = styles.bullet_style
        _apply_para_format(p, styles.bullet_para_fmt)
        run = p.add_run(text)
        _apply_run_format(run, styles.bullet_run_fmt)
        return p

    # ══════════════════════════════════════════
    # BUILD THE RESUME
    # ══════════════════════════════════════════

    # --- 1. HEADER ---
    add_para(
        optimized_data.get("name", "Your Name"),
        styles.name_para_fmt, styles.name_run_fmt
    )
    add_para(
        optimized_data.get("subtitle", ""),
        styles.subtitle_para_fmt, styles.subtitle_run_fmt
    )
    
    # Contact line — single run is fine, or split if links provided
    contact = optimized_data.get("contact_line", "")
    add_para(contact, styles.contact_para_fmt, styles.contact_run_fmt)

    # Spacer
    add_para("", styles.spacer_para_fmt, None)

    # --- 2. PROFESSIONAL SUMMARY ---
    add_section_heading("PROFESSIONAL SUMMARY")
    add_para(
        optimized_data.get("summary", ""),
        styles.body_para_fmt, styles.body_run_fmt
    )

    # --- 3. TECHNICAL SKILLS (table) ---
    add_section_heading("TECHNICAL SKILLS")

    skills = optimized_data.get("skills", [])
    if skills:
        # Create a table with same structure as template
        table = doc.add_table(rows=len(skills), cols=2)
        # Copy table properties from template (borders, widths)
        tblPr = styles.table_xml.find(qn('w:tblPr'))
        if tblPr is not None:
            existing_tblPr = table._element.find(qn('w:tblPr'))
            if existing_tblPr is not None:
                table._element.remove(existing_tblPr)
            table._element.insert(0, copy.deepcopy(tblPr))

        # Copy column grid from template
        tblGrid = styles.table_xml.find(qn('w:tblGrid'))
        if tblGrid is not None:
            existing_grid = table._element.find(qn('w:tblGrid'))
            if existing_grid is not None:
                table._element.remove(existing_grid)
            table._element.insert(1, copy.deepcopy(tblGrid))

        for i, skill in enumerate(skills):
            # Category cell
            cat_cell = table.rows[i].cells[0]
            cat_para = cat_cell.paragraphs[0]
            _apply_para_format(cat_para, styles.skill_cat_para_fmt)
            run = cat_para.add_run(skill.get("category", ""))
            _apply_run_format(run, styles.skill_cat_run_fmt)

            # Copy cell properties from template row for styling
            template_row = styles.table_xml.findall(qn('w:tr'))
            if i < len(template_row):
                # Copy cell properties (width, borders, shading)
                src_cells = template_row[min(i, len(template_row)-1)].findall(qn('w:tc'))
                if len(src_cells) >= 2:
                    src_tcPr = src_cells[0].find(qn('w:tcPr'))
                    if src_tcPr is not None:
                        existing = cat_cell._element.find(qn('w:tcPr'))
                        if existing is not None:
                            cat_cell._element.remove(existing)
                        cat_cell._element.insert(0, copy.deepcopy(src_tcPr))

                    src_tcPr2 = src_cells[1].find(qn('w:tcPr'))
                    if src_tcPr2 is not None:
                        val_cell = table.rows[i].cells[1]
                        existing2 = val_cell._element.find(qn('w:tcPr'))
                        if existing2 is not None:
                            val_cell._element.remove(existing2)
                        val_cell._element.insert(0, copy.deepcopy(src_tcPr2))

            # Skills cell
            val_cell = table.rows[i].cells[1]
            val_para = val_cell.paragraphs[0]
            _apply_para_format(val_para, styles.skill_val_para_fmt)
            run = val_para.add_run(skill.get("items", ""))
            _apply_run_format(run, styles.skill_val_run_fmt)

    # --- 4. WORK EXPERIENCE ---
    add_section_heading("WORK EXPERIENCE")

    for job in optimized_data.get("experience", []):
        # Job title + dates
        add_multi_run_para([
            (job.get("title", ""), styles.job_title_run_fmt),
            ("\t" + job.get("dates", ""), styles.job_dates_run_fmt),
        ], styles.job_title_para_fmt)

        # Company + location
        add_multi_run_para([
            (job.get("company", ""), styles.company_run_fmt),
            ("  |  " + job.get("location", ""), styles.company_loc_run_fmt),
        ], styles.company_para_fmt)

        # Bullet points
        for bullet in job.get("bullets", []):
            add_bullet(bullet)

    # --- 5. KEY PROJECTS ---
    projects = optimized_data.get("projects", [])
    if projects:
        add_section_heading("KEY PROJECTS")
        for proj in projects:
            # Project name + tech stack
            add_multi_run_para([
                (proj.get("name", "") + "  ", styles.project_title_run_fmt),
                ("| " + proj.get("tech", ""), styles.project_tech_run_fmt),
            ], styles.project_title_para_fmt)

            for bullet in proj.get("bullets", []):
                add_bullet(bullet)

    # --- 6. EDUCATION ---
    add_section_heading("EDUCATION")
    for edu in optimized_data.get("education", []):
        add_multi_run_para([
            (edu.get("degree", "") + "  —  " + edu.get("institution", ""), styles.edu_run_fmt),
            ("\t" + edu.get("dates", ""), styles.edu_dates_run_fmt),
        ], styles.edu_para_fmt)

        coursework = edu.get("coursework", "")
        if coursework:
            add_para(coursework, styles.coursework_para_fmt, styles.coursework_run_fmt)

    # --- 7. CERTIFICATIONS ---
    certs = optimized_data.get("certifications", [])
    if certs:
        add_section_heading("CERTIFICATIONS")
        for cert in certs:
            add_bullet(cert)

    # --- 8. ADDITIONAL ---
    additional = optimized_data.get("additional", [])
    if additional:
        add_section_heading("ADDITIONAL")
        for item in additional:
            add_bullet(item)

    # ── Save to BytesIO ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def parse_llm_json(llm_response: str) -> dict:
    """Extract JSON from the LLM response, handling markdown code fences.
    
    The LLM might wrap JSON in ```json ... ``` or return it with extra text.
    This function robustly extracts the JSON object.
    """
    # Try to find JSON in code fence
    fence_match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', llm_response, re.DOTALL)
    if fence_match:
        json_str = fence_match.group(1).strip()
    else:
        # Try to find raw JSON object
        brace_start = llm_response.find('{')
        brace_end = llm_response.rfind('}')
        if brace_start != -1 and brace_end != -1:
            json_str = llm_response[brace_start:brace_end + 1]
        else:
            raise ValueError("Could not find JSON in LLM response")

    return json.loads(json_str)
